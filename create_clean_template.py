"""
Copy title page + zavdannya pages (P0-P82) exactly as-is from templates.
No data filling, no modifications whatsoever.
"""
import sys, os, copy
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from lxml import etree

TITLE_PATH = r"C:\Users\viach\Desktop\word_data\data\практика\ДР Практика_Титулка (1).docx"
ZAVD_PATH  = r"C:\Users\viach\Desktop\word_data\data\практика\ШАБЛОН_Звіту_з_практики_+_зразок_змісту (2).docx"
OUT_DIR    = r"C:\Users\viach\Desktop\word_data\data\практика"
OUTPUT     = os.path.join(OUT_DIR, "Звіт_чистий_шаблон.docx")

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def insert_before_sectpr(body, element):
    """Insert element before body's sectPr (last child must be sectPr in valid OOXML)."""
    sectpr = body.find('w:sectPr', NSMAP)
    if sectpr is not None:
        sectpr.addprevious(element)
    else:
        body.append(element)

# Step 1: Open title template as base (exact copy)
print("Opening title template as base...")
doc = Document(TITLE_PATH)
body = doc.element.body

# Step 2: Add section break after title page to separate it from zavdannya
# Copy the title page's section properties and insert as paragraph-level sectPr
title_sectpr = body.find('w:sectPr', NSMAP)
title_sectpr_copy = copy.deepcopy(title_sectpr)

# Create a paragraph with section break (next page) using title's page settings
break_para = etree.SubElement(body, f'{{{NSMAP["w"]}}}p')
ppr = etree.SubElement(break_para, f'{{{NSMAP["w"]}}}pPr')
ppr.append(title_sectpr_copy)
# Set section break type to next page
sect_type = etree.SubElement(title_sectpr_copy, f'{{{NSMAP["w"]}}}type')
sect_type.set(f'{{{NSMAP["w"]}}}val', 'nextPage')

# Move this break paragraph before body's sectPr
body.remove(break_para)
insert_before_sectpr(body, break_para)

print(f"  Title page: {len(doc.paragraphs)} paragraphs, kept as-is.")

# Step 3: Copy zavdannya paragraphs P0-P82 as-is
print("Copying Завдання pages (P0-P82) from template...")
doc_z = Document(ZAVD_PATH)
zav_body = doc_z.element.body

zav_elements = []
for child in list(zav_body):
    tag = etree.QName(child.tag).localname
    if tag == 'sectPr':
        continue
    zav_elements.append(child)

# P0-P82 = first 83 paragraph-level elements
# But we need to count only <w:p> elements to match paragraph indices
p_count = 0
for elem in zav_elements:
    tag = etree.QName(elem.tag).localname
    cloned = copy.deepcopy(elem)
    insert_before_sectpr(body, cloned)
    if tag == 'p':
        p_count += 1
        if p_count > 82:
            break
    elif tag == 'tbl':
        pass  # tables between paragraphs — copy them too

print(f"  Copied {p_count} paragraphs from Завдання template.")

# Step 4: Update body's sectPr to match zavdannya page settings
# (the last section inherits zavdannya's page layout)
zav_sectpr = doc_z.element.body.find('w:sectPr', NSMAP)
if zav_sectpr is not None:
    old_sectpr = body.find('w:sectPr', NSMAP)
    new_sectpr = copy.deepcopy(zav_sectpr)
    # Remove header/footer references from zavdannya's sectPr to avoid missing rel errors
    for ref_tag in ['w:headerReference', 'w:footerReference']:
        for ref in new_sectpr.findall(ref_tag, NSMAP):
            new_sectpr.remove(ref)
    if old_sectpr is not None:
        body.replace(old_sectpr, new_sectpr)
    else:
        body.append(new_sectpr)

# Step 5: Make sure NO footers have page numbers on any section
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        for run in p.runs:
            run.text = ""
        # Also clear any field codes (PAGE fields)
        for elem in list(p._element):
            tag = etree.QName(elem.tag).localname
            if tag == 'r':
                has_field = elem.find('.//w:fldChar', NSMAP) is not None or elem.find('.//w:instrText', NSMAP) is not None
                if has_field:
                    p._element.remove(elem)

print("  Cleared all footers (no page numbering).")

# Save
doc.save(OUTPUT)
print(f"\nSaved: {OUTPUT}")
print(f"Size: {os.path.getsize(OUTPUT) / 1024:.1f} KB")
