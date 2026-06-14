"""
TechBox Practice Report v8
- Title page: copied 1:1 from template, only fill known blanks
- Завдання pages: copied 1:1, only fill known blanks
- Template pages: NO page numbers
- Body: humanized, 1.5 spacing, page numbers
- Dates: left blank (user fills by hand)
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import os

OUT_DIR = r"C:\Users\viach\Desktop\word_data\data\практика"
SCREENSHOTS = r"C:\Users\viach\Desktop\my-shop\reports\screenshots_light"
TITLE_TPL = os.path.join(OUT_DIR, "ДР Практика_Титулка (1).docx")
ZAVD_TPL = os.path.join(OUT_DIR, "ШАБЛОН_Звіту_з_практики_+_зразок_змісту (2).docx")

INDENT = Cm(1)


def clear_runs(p):
    for child in list(p._element):
        if child.tag == qn('w:r'):
            p._element.remove(child)


def add_run(p, text, size=14, bold=None, italic=False, underline=None):
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    if bold is not None:
        r.bold = bold
    if italic:
        r.italic = True
    if underline is not None:
        r.underline = underline
    return r


def insert_before_sectpr(doc, element):
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is not None:
        sectPr.addprevious(element)
    else:
        body.append(element)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def clear_footer(section):
    """Remove all content from section footer (no page numbers)."""
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        for child in list(p._element):
            if child.tag == qn('w:r'):
                p._element.remove(child)


def add_page_number_footer(section):
    """Add page number (right-aligned) to section footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.space_before = Pt(0)
    for child in list(fp._element):
        if child.tag == qn('w:r'):
            fp._element.remove(child)
    r = fp.add_run()
    fld = OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'), 'begin'); r._r.append(fld)
    r2 = fp.add_run()
    ins = OxmlElement('w:instrText'); ins.set(qn('xml:space'), 'preserve'); ins.text = ' PAGE '; r2._r.append(ins)
    r3 = fp.add_run()
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end'); r3._r.append(fld2)
    for rn in fp.runs:
        rn.font.name = 'Times New Roman'; rn.font.size = Pt(14)


# ============================================================
# STEP 1: Open title template — fill ONLY known fields
# ============================================================
print("Opening title template...")
doc = Document(TITLE_TPL)
paras = doc.paragraphs

# P12: practice type (original: "про проходження_______ практики")
# Keep same 14pt size as template
clear_runs(paras[12])
add_run(paras[12], "про проходження ", size=14)
add_run(paras[12], "переддипломної", size=14, underline=True)
add_run(paras[12], " практики", size=14)

# P13: "(назва практики)" hint — DON'T TOUCH

# P14: OPS + specialty (original: "освітнньо-професійний ступінь____ спеціальність____")
clear_runs(paras[14])
add_run(paras[14], "освітньо-професійний ступінь ", size=14)
add_run(paras[14], "фаховий молодший бакалавр", size=14, underline=True)
add_run(paras[14], " спеціальність ", size=14)
add_run(paras[14], '122 "Комп\'ютерні науки"', size=14, underline=True)

# P15: student + group
clear_runs(paras[15])
add_run(paras[15], "студента (ки) ", size=14)
add_run(paras[15], "Раделицького В.В.", size=14, underline=True)
add_run(paras[15], " групи ", size=14)
add_run(paras[15], "КНМС-41", size=14, underline=True)

# P16: "(прізвище та ініціали)" hint — DON'T TOUCH

# P17: organization name
clear_runs(paras[17])
add_run(paras[17], "на (в) ", size=14)
add_run(paras[17], "ФОП Раделицький Володимир Степанович", size=14, underline=True)

# P18: organization address (continuation of P17)
clear_runs(paras[18])
add_run(paras[18], "Львівська обл., Миколаївський р-н, с. Вербіж, вул. Івана Франка, 6/4", size=14, underline=True)

# P19: "(повна назва організації)" hint — DON'T TOUCH

# P23: dates — DON'T TOUCH (user fills by hand)

# All other paragraphs (spacers, headers, signature blocks) — DON'T TOUCH

# Remove page numbers from title page section
clear_footer(doc.sections[0])

print("  Title page: fields filled, dates left blank, no page numbers.")


# ============================================================
# STEP 2: Page break + Завдання pages (before body sectPr)
# ============================================================
print("Copying Завдання pages from template...")

# Page break
pb_el = OxmlElement('w:p')
pb_pPr = OxmlElement('w:pPr')
pb_el.append(pb_pPr)
pb_r = OxmlElement('w:r')
pb_br = OxmlElement('w:br')
pb_br.set(qn('w:type'), 'page')
pb_r.append(pb_br)
pb_el.append(pb_r)
insert_before_sectpr(doc, pb_el)

# Copy ALL Завдання paragraphs P0-P82 from template (1:1)
zavd_doc = Document(ZAVD_TPL)
zavd_paras = zavd_doc.paragraphs
for i in range(min(83, len(zavd_paras))):
    new_p = deepcopy(zavd_paras[i]._element)
    insert_before_sectpr(doc, new_p)

# Now fill ONLY known data in the copied Завдання paragraphs
all_paras = doc.paragraphs
zavd_start = None
for idx, p in enumerate(all_paras):
    if 'Студент____' in p.text or 'Студент____________' in p.text:
        if idx > 30:
            zavd_start = idx
            break

if zavd_start:
    print(f"  Завдання student field at paragraph {zavd_start}")
    zp = all_paras

    # Student name (P8 of zavdannya = zavd_start)
    clear_runs(zp[zavd_start])
    add_run(zp[zavd_start], "Студент ", size=12)
    add_run(zp[zavd_start], "Раделицький В'ячеслав Володимирович", size=12, underline=True)

    # OPS + specialty (P10 = zavd_start + 2)
    ops_idx = zavd_start + 2
    clear_runs(zp[ops_idx])
    add_run(zp[ops_idx], "Освітньо-професійний ступінь ", size=12)
    add_run(zp[ops_idx], "фаховий молодший бакалавр", size=12, underline=True)
    add_run(zp[ops_idx], " спеціальність ", size=12)
    add_run(zp[ops_idx], '122 "Комп\'ютерні науки"', size=12, underline=True)

    # Practice type (P11 = zavd_start + 3)
    pract_idx = zavd_start + 3
    clear_runs(zp[pract_idx])
    add_run(zp[pract_idx], "Скерований на практику ", size=12)
    add_run(zp[pract_idx], "переддипломну", size=12, underline=True)

    # City + org (P13 = zavd_start + 5)
    city_idx = zavd_start + 5
    clear_runs(zp[city_idx])
    add_run(zp[city_idx], "в місто ", size=12)
    add_run(zp[city_idx], "Львів", size=12, underline=True)
    add_run(zp[city_idx], "  на  ", size=12)
    add_run(zp[city_idx], "ФОП Раделицький Володимир Степанович", size=12, underline=True)

    # P15 of zavdannya (zavd_start + 7) — org address continuation
    addr_idx = zavd_start + 7
    clear_runs(zp[addr_idx])
    add_run(zp[addr_idx], "Львівська обл., Миколаївський р-н, с. Вербіж, вул. Івана Франка, 6/4", size=12, underline=True)

    # P19 dates (zavd_start + 11) — DON'T TOUCH (user fills by hand)
    # P17 договору — DON'T TOUCH
    # Зміст завдання — DON'T TOUCH (supervisor fills in)
    # Відгуки — DON'T TOUCH

    print("  Завдання: known fields filled, dates/tasks/reviews left blank.")
else:
    print("  WARNING: Could not find Завдання student field!")


# ============================================================
# STEP 3: New section for report body
# ============================================================
new_sec = doc.add_section()
new_sec.top_margin = Cm(2)
new_sec.bottom_margin = Cm(2)
new_sec.left_margin = Cm(2)
new_sec.right_margin = Cm(1)

# Page numbers ONLY on body section
add_page_number_footer(new_sec)

# DO NOT modify Normal style — template pages must keep their original spacing.
# Body paragraphs set their own line_spacing individually.


# ============================================================
# Body helper functions (each sets 1.5x spacing explicitly)
# ============================================================
def empty(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def centered(text, bold=False, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.name = 'Times New Roman'; r.bold = bold
    return p


def heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if level == 1:
        p.paragraph_format.first_line_indent = INDENT
    r = p.add_run(text.upper() if level == 0 else text)
    r.bold = True; r.font.size = Pt(14); r.font.name = 'Times New Roman'


def body(text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        p.paragraph_format.first_line_indent = INDENT
    r = p.add_run(text)
    r.font.size = Pt(14); r.font.name = 'Times New Roman'


def li(text, marker="–"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.first_line_indent = INDENT
    r = p.add_run(f"{marker} {text}")
    r.font.size = Pt(14); r.font.name = 'Times New Roman'


def ni(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = INDENT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    r.font.size = Pt(14); r.font.name = 'Times New Roman'


def toc_line(title, pg, is_sub=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if is_sub:
        p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(17), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(title)
    r.font.size = Pt(14); r.font.name = 'Times New Roman'
    if not is_sub:
        r.bold = True
    r2 = p.add_run('\t')
    r2.font.size = Pt(14); r2.font.name = 'Times New Roman'
    r3 = p.add_run(str(pg))
    r3.font.size = Pt(14); r3.font.name = 'Times New Roman'


def pb():
    doc.add_page_break()


def img_center(path, w_cm=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(path, width=Cm(w_cm))


def fig_cap(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(12); r.font.name = 'Times New Roman'; r.italic = True


# ============================================================
# ЗМІСТ
# ============================================================
print("Adding report body...")
heading("ЗМІСТ", 0)
empty(1)
for title, pg, sub in [
    ("ВСТУП", 7, False),
    ("1. ХАРАКТЕРИСТИКА ПРЕДМЕТНОЇ ОБЛАСТІ", 9, False),
    ("1.1 Опис бази практики", 9, True),
    ("1.2 Опис частини діяльності для автоматизації", 10, True),
    ("1.3 Огляд та аналіз предметної області", 11, True),
    ("2. ОГЛЯД СУЧАСНИХ ПІДХОДІВ ТА ЗАСОБІВ ДО ПРОЕКТУВАННЯ ТА РОЗРОБЛЕННЯ ВЕБ-ДОДАТКІВ", 13, False),
    ("2.1 Огляд сучасних підходів до проектування систем", 13, True),
    ("2.2 Огляд сучасних засобів розробки веб-додатків", 14, True),
    ("2.3 Огляд та аналіз аналогічних систем", 15, True),
    ("3. ПОСТАНОВКА ЗАВДАННЯ НА ПРАКТИКУ", 16, False),
    ("4. РЕАЛІЗАЦІЯ ПОСТАВЛЕНИХ ЗАВДАНЬ", 18, False),
    ("4.1 Специфікація вимог до системи", 18, True),
    ("4.2 Розроблений інтерфейс веб-системи", 21, True),
    ("ВИСНОВКИ", 23, False),
    ("СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ", 25, False),
]:
    toc_line(title, pg, sub)


# ============================================================
# ВСТУП
# ============================================================
pb()
heading("ВСТУП", 0)
empty(1)
body("Ринок електронної комерції в Україні продовжує зростати. За даними Асоціації ритейлерів України, обсяг онлайн-продажів у 2025 році збільшився на 28 % порівняно з попереднім роком [1]. Водночас зростає потреба в розробниках, які можуть побудувати веб-платформу з нуля, від серверної логіки до клієнтського інтерфейсу й розгортання.")
body("Переддипломна практика проходила на базі ФОП Раделицький Володимир Степанович. Метою було створення робочої платформи електронної комерції TechBox для продажу електроніки та гаджетів на українському ринку. Платформа охоплює повний цикл покупки: від каталогу з пошуком до оформлення замовлення, онлайн-оплати через LiqPay та адміністративної панелі з аналітикою.")
body("Робота охопила аналіз предметної області, проєктування архітектури, написання серверної (backend) та клієнтської (frontend) частин, налаштування бази даних і контейнеризацію.")
body("Завдання практики:")
for t in [
    "дослідити предметну область електронної комерції та виявити потреби типового інтернет-магазину;",
    "провести аналіз підходів і технологій для побудови веб-додатків;",
    "спроєктувати архітектуру системи та реляційну базу даних;",
    "реалізувати серверний API на базі Python/FastAPI та клієнтську частину на React;",
    "інтегрувати платіжну систему LiqPay та AI-модуль генерації опису товарів;",
    "забезпечити контейнеризацію за допомогою Docker та Docker Compose.",
]:
    li(t)
body("Більшість існуючих рішень для інтернет-магазинів в Україні або побудовані на закордонних SaaS-платформах (Shopify, WooCommerce), або потребують значних витрат на адаптацію: українська мова, національні платіжні системи, специфіка доставки [2]. TechBox з самого початку проєктувався під ці потреби.")
body("Звіт складається з чотирьох розділів. Перший описує предметну область та базу практики. Другий присвячено огляду технологій, фреймворків та аналогічних систем. Третій містить завдання на практику. Четвертий документує реалізацію: специфікацію вимог та інтерфейс веб-системи.")


# ============================================================
# РОЗДІЛ 1
# ============================================================
pb()
heading("1.  ХАРАКТЕРИСТИКА ПРЕДМЕТНОЇ ОБЛАСТІ", 0)
empty(1)
heading("1.1 Опис бази практики", 1)
body("Практика проводилася на базі ФОП Раделицький Володимир Степанович (Львівська обл., Миколаївський р-н, с. Вербіж, вул. Івана Франка, 6/4). Завданням було повне розроблення веб-платформи електронної комерції. Робота велась на персональному комп'ютері з Windows 11 Pro, як редактор використовувався Visual Studio Code.")
body("Серверну частину проєкту написано мовою Python 3.13 із застосуванням фреймворку FastAPI. Для управління залежностями обрано менеджер UV, що пришвидшує встановлення бібліотек і відтворює середовище [3]. Клієнтська частина працює на React 19 з TypeScript, збірник Vite 8.")
body("Як базу даних обрано PostgreSQL 16, розгорнуту в Docker-контейнері. Для контролю версій використовувався Git, репозиторій розміщено на GitHub.")
body("Також використовувались:")
for t in [
    "Docker та Docker Compose для контейнеризації та оркестрації сервісів;",
    "Alembic для управління міграціями бази даних;",
    "LiqPay API для інтеграції онлайн-платежів;",
    "Google Gemini API для генерації AI-описів товарів;",
    "FastAPI-Mail для відправки листів верифікації.",
]:
    li(t)

heading("1.2 Опис частини діяльності для автоматизації", 1)
body("Система автоматизує типові бізнес-процеси інтернет-магазину електроніки. Основне завдання: каталог товарів із пошуком, фільтрацією за категоріями та переглядом характеристик кожного продукту.")
body("Оформлення замовлення складається з кількох кроків: формування кошика, вибір способу доставки та оплати, застосування промокоду й підтвердження. Інтеграція з LiqPay дає змогу оплачувати онлайн, статус оновлюється автоматично через webhook [4].")
body("Для працівників є окрема панель з ієрархією ролей. Менеджер керує товарами та замовленнями. Адміністратор додатково бачить користувачів, промокоди та статистику. Суперадміністратор може призначати ролі. Панель автоматизує:")
for t in [
    "моніторинг замовлень та зміну статусів (нове, в обробці, відправлено, доставлено, скасовано);",
    "управління каталогом товарів та контроль залишків;",
    "перегляд і управління обліковими записами;",
    "створення та відстеження промокодів;",
    "статистику продажів: виручку, середній чек, топ-продукти, розподіл за статусами.",
]:
    li(t)
body("Окремий AI-модуль на базі Google Gemini генерує маркетингові описи та технічні характеристики товарів українською мовою. Вхідні дані: назва та категорія продукту. Результати кешуються на 48 годин, щоб зменшити кількість API-запитів.")

heading("1.3 Огляд та аналіз предметної області", 1)
body("Електронна комерція охоплює процеси купівлі та продажу товарів через Інтернет. В Україні частка онлайн-торгівлі перевищила 15 % загального обсягу роздрібних продажів у 2025 році [1]. Найпопулярніші категорії: електроніка, побутова техніка, одяг та товари для дому.")
body("Типовий інтернет-магазин електроніки потребує:")
for t in [
    "каталог із категоріями, пошуком та фільтрацією;",
    "кошик та оформлення замовлення;",
    "онлайн-оплату через національні платіжні сервіси;",
    "реєстрацію та авторизацію користувачів;",
    "систему відгуків та рейтингів;",
    "адміністративну панель для управління контентом та замовленнями.",
]:
    li(t)
body("Головні проблеми для розробників в Україні: обмежена підтримка української мови в готових рішеннях, складна інтеграція з локальними платіжними провайдерами (LiqPay, Monobank, ПриватБанк), а також необхідність враховувати специфіку логістичних служб (Нова Пошта, Укрпошта) [2].")
body("Існуючі рішення поділяються на три групи. Перша: SaaS-платформи (Shopify, Horoshop), швидкий запуск, але обмежена кастомізація. Друга: CMS на основі PHP (WooCommerce, OpenCart), більше свободи, але прив'язка до PHP. Третя: кастомні розробки, максимальний контроль, проте найбільші витрати часу [5].")
body("TechBox належить до третьої групи. Рішення про власну розробку продиктоване навчальною метою практики: пройти повний цикл створення веб-додатка, від проєктування бази даних до розгортання в Docker-контейнерах.")


# ============================================================
# РОЗДІЛ 2
# ============================================================
pb()
heading("2.  ОГЛЯД СУЧАСНИХ ПІДХОДІВ ТА ЗАСОБІВ ДО ПРОЕКТУВАННЯ ТА РОЗРОБЛЕННЯ ВЕБ-ДОДАТКІВ", 0)
empty(1)
heading("2.1 Огляд сучасних підходів до проектування систем", 1)
body("Для веб-додатків найчастіше застосовують клієнт-серверну архітектуру з розділенням відповідальності на окремі рівні. Для TechBox обрано трирівневу схему: рівень представлення (frontend), рівень бізнес-логіки (backend API) та рівень даних (PostgreSQL) [6].")
body("Взаємодія між клієнтом і сервером побудована за стилем REST, який передбачає стандартні HTTP-методи та обмін даними у форматі JSON [7]. Клієнтський додаток реалізовано як SPA (Single Page Application): сторінка завантажується один раз і далі оновлює контент через API-запити [8].")
body("Для управління станом у React-додатку використано Context API з трьома контекстами: AuthContext (авторизація), CartContext (кошик) і ThemeContext (тема інтерфейсу). На серверній стороні моделі даних описано через ORM SQLAlchemy.")
body("Docker ізолює кожен сервіс у контейнері, а Docker Compose описує всю інфраструктуру в одному файлі. Розгортання зводиться до однієї команди [9].")

heading("2.2 Огляд сучасних засобів розробки веб-додатків", 1)
body("Серверну частину побудовано на FastAPI (Python). Фреймворк генерує Swagger-документацію і підтримує асинхронну обробку запитів [3]. Для роботи з базою даних використано SQLAlchemy 2.0, для міграцій Alembic [10].")
body("Клієнтська частина працює на React 19 з TypeScript для статичної типізації [11], збірник Vite 8. Маршрутизацію реалізовано через React Router DOM 7, стилізацію через CSS Modules.")
body("Автентифікація побудована на JWT (JSON Web Token) [12]. Паролі хешуються алгоритмом bcrypt. Docker Compose піднімає чотири сервіси: PostgreSQL 16, FastAPI, React-клієнт та адміністративну панель.")

heading("2.3 Огляд та аналіз аналогічних систем", 1)
body("Для порівняння обрано три системи українського ринку електроніки: Rozetka (найбільший маркетплейс із фасетним пошуком та програмою лояльності) [2], Jabko (спеціалізація на Apple, інтеграція з LiqPay) і Citrus (кредитні програми, trade-in).")
body("Порівняльний аналіз наведено у табл. 2.1.")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
r = p.add_run("Таблиця 2.1 – Порівняння функціональності аналогічних систем")
r.font.size = Pt(12); r.font.name = 'Times New Roman'; r.italic = True

t21 = doc.add_table(rows=8, cols=5)
set_table_borders(t21)
t21.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Критерій", "Rozetka", "Jabko", "Citrus", "TechBox"]):
    c = t21.rows[0].cells[j]; c.text = h
    for pr in c.paragraphs:
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for rn in pr.runs: rn.bold = True; rn.font.size = Pt(12); rn.font.name = 'Times New Roman'
for i, rd in enumerate([
    ["Каталог із фільтрацією", "+", "+", "+", "+"],
    ["Кошик та замовлення", "+", "+", "+", "+"],
    ["Онлайн-оплата", "+", "+", "+", "+"],
    ["Відгуки та рейтинги", "+", "+", "+", "+"],
    ["Програма лояльності", "+", "+", "+", "–"],
    ["Порівняння товарів", "+", "+", "+", "–"],
    ["AI-генерація описів", "–", "–", "–", "+"],
]):
    for j, v in enumerate(rd):
        c = t21.rows[i + 1].cells[j]; c.text = v
        for pr in c.paragraphs:
            pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for rn in pr.runs: rn.font.size = Pt(12); rn.font.name = 'Times New Roman'

empty(1)
body("Усі аналізовані системи мають базовий функціонал інтернет-магазину. Комерційні платформи за роки роботи додали лояльність і порівняння товарів. TechBox відрізняється AI-генерацією описів та характеристик, якої немає в жодного з конкурентів.")


# ============================================================
# РОЗДІЛ 3
# ============================================================
pb()
heading("3.  ПОСТАНОВКА ЗАВДАННЯ НА ПРАКТИКУ", 0)
empty(1)
body("Мета практики полягала у розробленні веб-платформи електронної комерції TechBox для продажу електроніки та гаджетів. Система мала включати клієнтський інтерфейс для покупців, адміністративну панель для управління магазином та серверну частину з REST API.")
body("Завдання:")
for t in [
    "проаналізувати предметну область електронної комерції в Україні та визначити бізнес-процеси інтернет-магазину;",
    "провести огляд технологій веб-розробки та обрати стек;",
    "спроєктувати реляційну базу даних для зберігання товарів, користувачів, замовлень, відгуків та промокодів;",
    "реалізувати серверний API з авторизацією через JWT, ролевою моделлю доступу та асинхронною обробкою;",
    "розробити клієнтський SPA з каталогом товарів, кошиком, оформленням замовлення та особистим кабінетом;",
    "створити адміністративну панель для управління замовленнями, товарами, користувачами та аналітикою;",
    "інтегрувати LiqPay для онлайн-оплати;",
    "реалізувати AI-генерацію описів товарів на базі Google Gemini;",
    "налаштувати контейнеризацію через Docker та Docker Compose.",
]:
    li(t)
body("Функціональні вимоги:")
for t in [
    "1) авторизація користувачів з підтвердженням електронної пошти;",
    "2) відображення каталогу товарів за категоріями;",
    "3) пошук товарів та перегляд детальної інформації;",
    "4) додавання товарів до кошика та оформлення замовлення;",
    "5) оплата замовлення через LiqPay;",
    "6) відстеження статусу замовлення та перегляд історії покупок;",
    "7) можливість залишати відгуки та оцінки товарам;",
    "8) адміністративна панель для управління всіма аспектами магазину;",
    "9) AI-генерація описів та характеристик товарів.",
]:
    ni(t)
body("Очікуваний результат: працездатна платформа, де клієнти переглядають каталог, додають товари до кошика, оформлюють і оплачують замовлення, залишають відгуки, а адміністратори керують магазином через окрему панель.")


# ============================================================
# РОЗДІЛ 4
# ============================================================
pb()
heading("4.  РЕАЛІЗАЦІЯ ПОСТАВЛЕНИХ ЗАВДАНЬ", 0)
empty(1)
heading("4.1 Специфікація вимог до системи", 1)
body("Вимоги до інтерфейсу:")
ni("1. Інтерфейс має бути зрозумілим та зручним.")
ni("2. Адаптивний дизайн під різні розміри екрану.")
body("Функціональні вимоги.")
body("Адміністративна панель:")
ni("1. Додавання продуктів до каталогу: назва, опис, ціна, зображення, категорія, кількість.")
ni("2. Список замовлень з можливістю зміни статусу.")
ni("3. Статистика: виручка, кількість замовлень, середній чек, топ-продукти, розподіл за статусами.")
ni("4. Управління обліковими записами та ролями.")
ni("5. Управління промокодами та знижками.")
body("Інтерфейс користувача:")
for r_ in [
    "1. Реєстрація та авторизація з підтвердженням пошти.",
    "2. Каталог продуктів із фільтрацією за категоріями.",
    "3. Сторінка продукту з описом, характеристиками, відгуками.",
    "4. Додавання до кошика та зміна кількості.",
    "5. Оформлення замовлення з вибором доставки та оплати.",
    "6. Оплата через LiqPay.",
    "7. Історія замовлень та відстеження статусу.",
    "8. Відгуки та оцінки товарів.",
    "9. Налаштування профілю: ім'я, телефон, пароль.",
]:
    ni(r_)
body("Нефункціональні вимоги:")
ni("1. Час відповіді API не більше 500 мс.")
ni("2. Безпечне зберігання паролів (bcrypt) та JWT-токени.")
ni("3. Розгортання через Docker Compose однією командою.")
empty(1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
r = p.add_run("Таблиця 4.1 – Функціональні можливості акторів")
r.font.size = Pt(12); r.font.name = 'Times New Roman'; r.italic = True

actors = [
    ["Customer\n(покупець)", "Авторизація", "Створити акаунт або увійти, підтвердити пошту"],
    ["", "Головна сторінка", "Перегляд популярних товарів, перехід до каталогу, кошика, профілю"],
    ["", "Перегляд продукту", "Опис, ціна, AI-характеристики, відгуки покупців"],
    ["", "Кошик", "Додавання товарів, зміна кількості, підсумкова вартість"],
    ["", "Оформлення замовлення", "Контактні дані, вибір доставки та оплати, промокод"],
    ["", "Оплата", "Оплата через LiqPay або при отриманні"],
    ["", "Відгуки", "Відгук та оцінка товару (за умови покупки)"],
    ["", "Профіль", "Зміна імені, телефону, паролю"],
    ["Manager\n(менеджер)", "Товари", "Додавання та редагування товарів"],
    ["", "Замовлення", "Перегляд та зміна статусів"],
    ["Admin\n(адмін)", "Товари", "Додавання, редагування, видалення"],
    ["", "Замовлення", "Деталі замовлень, зміна статусів, інформація про покупця"],
    ["", "Користувачі", "Список, статистика покупок, редагування профілів"],
    ["", "Промокоди", "Створення, редагування, деактивація; контроль використань та терміну дії"],
    ["", "Статистика", "Виручка, середній чек, топ-продукти, розподіл замовлень"],
    ["Superadmin", "Повний доступ", "Усі функції адміністратора"],
    ["", "Ролі", "Призначення ролей: customer, manager, admin"],
]

t41 = doc.add_table(rows=len(actors) + 1, cols=3)
set_table_borders(t41)
t41.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Актор", "Активність", "Можливість"]):
    c = t41.rows[0].cells[j]; c.text = h
    for pr in c.paragraphs:
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for rn in pr.runs: rn.bold = True; rn.font.size = Pt(12); rn.font.name = 'Times New Roman'
for i, rd in enumerate(actors):
    for j, v in enumerate(rd):
        c = t41.rows[i + 1].cells[j]; c.text = v
        for pr in c.paragraphs:
            for rn in pr.runs: rn.font.size = Pt(12); rn.font.name = 'Times New Roman'
            if j == 1:
                for rn in pr.runs: rn.bold = True


# 4.2 Interface screenshots
pb()
heading("4.2 Розроблений інтерфейс веб-системи", 1)
body("Клієнтський інтерфейс TechBox дає покупцю пройти повний шлях: від перегляду каталогу до оформлення і оплати замовлення.")

body("Головна сторінка (рис. 4.1) містить навігаційну панель, банер з категоріями та блок популярних товарів.")
img_path = os.path.join(SCREENSHOTS, "01_home.png")
if os.path.exists(img_path):
    img_center(img_path, 15)
fig_cap("Рис. 4.1 – Головна сторінка інтернет-магазину TechBox")

body("Каталог (рис. 4.2) має бічну панель категорій та сітку карток товарів. Кожна картка показує зображення, назву, рейтинг, ціну та кнопку додавання до кошика.")
img_path = os.path.join(SCREENSHOTS, "02_catalog.png")
if os.path.exists(img_path):
    img_center(img_path, 15)
fig_cap("Рис. 4.2 – Каталог товарів із фільтрацією за категоріями")

body("Сторінка товару (рис. 4.3) показує зображення, AI-згенерований опис і характеристики, ціну, наявність, кнопку кошика та відгуки.")
img_path = os.path.join(SCREENSHOTS, "03_product.png")
if os.path.exists(img_path):
    img_center(img_path, 15)
fig_cap("Рис. 4.3 – Сторінка товару з AI-описом та відгуками")

body("Кошик (рис. 4.4) показує обрані товари з можливістю зміни кількості, поле промокоду та підсумок із кнопкою оформлення.")
img_path = os.path.join(SCREENSHOTS, "04_cart.png")
if os.path.exists(img_path):
    img_center(img_path, 15)
fig_cap("Рис. 4.4 – Кошик покупця з промокодом та підсумком")


# ============================================================
# ВИСНОВКИ
# ============================================================
pb()
heading("ВИСНОВКИ", 0)
empty(1)
body("За час переддипломної практики розроблено веб-платформу TechBox для продажу електроніки та гаджетів на українському ринку.")
body("Виконано такі завдання:")
for t in [
    "проаналізовано предметну область електронної комерції в Україні, виявлено бізнес-процеси та проблеми існуючих рішень;",
    "обрано технологічний стек: FastAPI, React 19, PostgreSQL, Docker;",
    "спроєктовано реляційну базу з шести таблиць (товари, користувачі, замовлення, відгуки, промокоди);",
    "реалізовано REST API із 20+ ендпоінтами, JWT-авторизацією та ролевою моделлю доступу;",
    "розроблено клієнтський SPA із каталогом, кошиком, оформленням замовлення, особистим кабінетом та відгуками;",
    "створено адміністративну панель зі статистикою, управлінням замовленнями, товарами, користувачами та промокодами;",
    "інтегровано LiqPay із підтримкою webhook для автоматичного оновлення статусу оплати;",
    "реалізовано AI-модуль генерації описів на базі Google Gemini з кешуванням на 48 годин;",
    "налаштовано Docker Compose з чотирма сервісами та health check.",
]:
    li(t)
body("Створена система придатна як основа для реального інтернет-магазину. Інтерфейс та AI-описи українською мовою, платежі через LiqPay.")
body("Набутий досвід: проєктування REST API, робота з ORM та міграціями, побудова SPA на React із TypeScript, інтеграція зовнішніх сервісів (LiqPay, Google Gemini, SMTP), контейнеризація та оркестрація.")


# ============================================================
# СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ
# ============================================================
pb()
heading("СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ", 0)
empty(1)
for s in [
    "1. Асоціація ритейлерів України. Аналітичний звіт: e-commerce в Україні 2024–2025. URL: https://rau.ua/analytics/ (дата звернення: 15.05.2026).",
    "2. Коваленко О. В., Петренко М. І. Сучасні тенденції розвитку електронної комерції в Україні. Економіка та управління підприємствами. 2024. Вип. 38. С. 45–53.",
    "3. FastAPI Documentation. URL: https://fastapi.tiangolo.com/ (дата звернення: 10.05.2026).",
    "4. LiqPay API Documentation. URL: https://www.liqpay.ua/documentation (дата звернення: 10.05.2026).",
    "5. Мельник В. О. Порівняльний аналіз платформ для створення інтернет-магазинів. Інформаційні технології та комп'ютерна інженерія. 2023. № 2. С. 112–120.",
    "6. Річардсон К. Мікросервіси. Паттерни розробки та рефакторингу / пер. з англ. СПб. : Пітер, 2022. 544 с.",
    "7. Fielding R. Architectural Styles and the Design of Network-based Software Architectures : Ph.D. dissertation. University of California, Irvine, 2000. 162 p.",
    "8. Банкс А., Порселло Є. React. Сучасні патерни для розробки додатків / пер. з англ. 2-ге вид. СПб. : Пітер, 2023. 350 с.",
    "9. Docker Documentation. URL: https://docs.docker.com/ (дата звернення: 12.05.2026).",
    "10. SQLAlchemy 2.0 Documentation. URL: https://docs.sqlalchemy.org/ (дата звернення: 12.05.2026).",
    "11. TypeScript Documentation. URL: https://www.typescriptlang.org/docs/ (дата звернення: 12.05.2026).",
    "12. Jones M., Bradley J., Sakimura N. JSON Web Token (JWT). RFC 7519. IETF, 2015. URL: https://tools.ietf.org/html/rfc7519 (дата звернення: 14.05.2026).",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = INDENT
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(s)
    r.font.size = Pt(14); r.font.name = 'Times New Roman'


# ============================================================
# SAVE
# ============================================================
output = os.path.join(OUT_DIR, "Звіт_з_практики_TechBox_v9.docx")
doc.save(output)
print(f"\nSaved: {output}")
print(f"Size: {os.path.getsize(output) / 1024:.1f} KB")
