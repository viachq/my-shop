from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


SOURCE = Path(r"C:\Users\viach\Desktop\Звіт_з_практики_TechBox_заповнити.docx")
WORKING = Path("doc_work/practice_fill.docx")

DATA = {
    "full_name": "Раделицький В'ячеслав Володимирович",
    "short_name": "Раделицький В. В.",
    "group": "КНМС-41",
    "degree": "фаховий молодший бакалавр",
    "specialty": "122 \"Комп'ютерні науки\"",
    "practice": "переддипломної",
    "city": "Львів",
    "org_full": "ФОП Раделицький Володимир Степанович",
    "org_full_with_short": "ФОП Раделицький Володимир Степанович (Раделицький В. С.)",
    "org_head_short": "Раделицький В. С.",
    "org_head_role": "ФОП Раделицький В. С.",
    "dept_head_full": "Машевська Марта Володимирівна",
    "dept_head_short": "Машевська М. В.",
}


def add_run(paragraph, text, *, size=14, bold=False, italic=False, underline=False):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    return run


def set_parts(doc, idx, parts, align=None):
    paragraph = doc.paragraphs[idx]
    paragraph.clear()
    if align is not None:
        paragraph.alignment = align

    for part in parts:
        add_run(
            paragraph,
            part.get("text", ""),
            size=part.get("size", 14),
            bold=part.get("bold", False),
            italic=part.get("italic", False),
            underline=part.get("underline", False),
        )


def patch_page_number_start(docx_path, start="1"):
    with ZipFile(docx_path, "r") as zin, NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
        with ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    xml = data.decode("utf-8")
                    xml = xml.replace('w:pgNumType w:start="4"', f'w:pgNumType w:start="{start}"')
                    data = xml.encode("utf-8")
                zout.writestr(item, data)
    tmp_path.replace(docx_path)


def main():
    WORKING.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(SOURCE))

    set_parts(
        doc,
        12,
        [
            {"text": "про проходження ", "size": 14},
            {"text": DATA["practice"], "size": 14, "underline": True},
            {"text": " практики", "size": 14},
        ],
    )
    set_parts(
        doc,
        14,
        [
            {"text": "освітньо-професійний ступінь ", "size": 10.5},
            {"text": DATA["degree"], "size": 10.5, "underline": True},
            {"text": " спеціальність ", "size": 10.5},
            {"text": DATA["specialty"], "size": 10.5, "underline": True},
        ],
    )
    set_parts(
        doc,
        15,
        [
            {"text": "студента ", "size": 14},
            {"text": DATA["short_name"], "size": 14, "underline": True},
            {"text": " групи ", "size": 14},
            {"text": DATA["group"], "size": 14, "underline": True},
        ],
    )
    set_parts(
        doc,
        17,
        [
            {"text": "на (в) ", "size": 13},
            {"text": DATA["org_full"], "size": 13, "underline": True},
        ],
    )
    set_parts(
        doc,
        18,
        [{"text": "(Раделицький В. С.)", "size": 13, "underline": True}],
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_parts(
        doc,
        23,
        [{"text": "від „20” квітня 2026 р.               до „16” травня 2026 р.", "size": 14}],
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_parts(
        doc,
        32,
        [
            {"text": "від організації ", "size": 14},
            {"text": DATA["org_head_role"], "size": 14, "underline": True},
        ],
    )
    set_parts(
        doc,
        39,
        [
            {"text": "Керівник практики від кафедри  ", "size": 14},
            {"text": DATA["dept_head_short"], "size": 14, "underline": True},
        ],
    )

    # The template relied on Word's cached page positions. Add a real break so
    # the second form never drifts onto the cover page after typed fields change.
    doc.paragraphs[44].add_run().add_break(WD_BREAK.PAGE)

    set_parts(
        doc,
        53,
        [
            {"text": "Студент ", "size": 12},
            {"text": DATA["full_name"], "size": 12, "underline": True},
        ],
    )
    set_parts(
        doc,
        55,
        [
            {"text": "Освітньо-професійний ступінь ", "size": 10.5},
            {"text": DATA["degree"], "size": 10.5, "underline": True},
            {"text": " спеціальність ", "size": 10.5},
            {"text": DATA["specialty"], "size": 10.5, "underline": True},
        ],
    )
    set_parts(
        doc,
        56,
        [
            {"text": "Скерований на практику ", "size": 12},
            {"text": DATA["practice"], "size": 12, "underline": True},
        ],
    )
    set_parts(
        doc,
        58,
        [
            {"text": "в місто ", "size": 12},
            {"text": DATA["city"], "size": 12, "underline": True},
            {"text": " на", "size": 12},
        ],
    )
    set_parts(
        doc,
        60,
        [{"text": DATA["org_full_with_short"], "size": 12, "underline": True}],
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_parts(
        doc,
        64,
        [
            {"text": "Термін практики: від ", "size": 12},
            {"text": "20.04.2026р.", "size": 12, "underline": True},
            {"text": " до ", "size": 12},
            {"text": "16.05.2026р.", "size": 12, "underline": True},
        ],
    )
    set_parts(
        doc,
        66,
        [
            {"text": "Керівник практики від кафедри ", "size": 12},
            {"text": DATA["dept_head_full"], "size": 12, "underline": True},
        ],
    )
    set_parts(
        doc,
        81,
        [{"text": "Прибув на базу практики \t\t„20” квітня 2026 року", "size": 12}],
    )
    set_parts(
        doc,
        82,
        [
            {"text": "______________________________", "size": 10},
            {"text": "\t\t\t", "size": 10},
            {"text": DATA["org_head_short"], "size": 12, "underline": True},
        ],
    )
    set_parts(
        doc,
        90,
        [{"text": "Вибув з бази практики  \t\t “16” травня 2026 року", "size": 12}],
    )
    set_parts(
        doc,
        91,
        [
            {"text": "______________________________", "size": 10},
            {"text": "\t\t\t", "size": 10},
            {"text": DATA["org_head_short"], "size": 12, "underline": True},
        ],
    )

    # Remove trailing empty section paragraphs that Word renders as a blank page.
    for paragraph in reversed(doc.paragraphs):
        if paragraph.text.strip():
            break
        paragraph._element.getparent().remove(paragraph._element)

    doc.save(str(WORKING))
    patch_page_number_start(WORKING)
    print(WORKING.resolve())


if __name__ == "__main__":
    main()
