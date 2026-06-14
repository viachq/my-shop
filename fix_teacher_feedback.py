"""Apply teacher feedback fixes to TechBox report."""
import sys, os, shutil, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
SRC = r"C:\Users\viach\Desktop\word_data\data\практика\Звіт_з_практики_TechBox_v9 - Copy.docx"
BAK = SRC.replace(".docx", "_backup4.docx")
if not os.path.exists(BAK):
    shutil.copy2(SRC, BAK)
    print("backup ->", BAK)

d = Document(SRC)
ps = d.paragraphs

def numid(p):
    ppr = p._p.find(W + 'pPr')
    if ppr is None: return None
    npr = ppr.find(W + 'numPr')
    if npr is None: return None
    n = npr.find(W + 'numId')
    return n.get(W + 'val') if n is not None else None

# ============================================================
# FIX 1: Figure & table captions — regular text (not italic, 14pt TNR)
# ============================================================
n_cap = 0
for p in ps:
    t = p.text.strip()
    if t.startswith('Рис.') or t.startswith('Таблиця'):
        for r in p.runs:
            r.italic = False
            r.font.size = Pt(14)
            r.font.name = 'Times New Roman'
        n_cap += 1
print(f"FIX 1: captions de-italicized & set 14pt: {n_cap}")

# ============================================================
# FIX 2: ВИСНОВКИ — rewrite as flowing paragraphs (no bullet list)
# ============================================================
# Find ВИСНОВКИ boundaries
vysn_start = None
vysn_end = None
for i, p in enumerate(ps):
    t = p.text.strip()
    if t == 'ВИСНОВКИ':
        vysn_start = i
    if vysn_start is not None and i > vysn_start:
        if t == 'СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ' or t.startswith('ДОДАТ'):
            vysn_end = i
            break

if vysn_start is not None and vysn_end is not None:
    # Delete all paragraphs between ВИСНОВКИ heading and next section
    # (indices vysn_start+1 to vysn_end-1)
    to_remove = []
    for i in range(vysn_start + 1, vysn_end):
        to_remove.append(ps[i])
    for p in to_remove:
        p._p.getparent().remove(p._p)
    print(f"FIX 2: removed {len(to_remove)} old ВИСНОВКИ paragraphs")

    # Now insert new flowing text AFTER the ВИСНОВКИ heading
    # We need to insert before the next remaining element (which is now СПИСОК...)
    vysn_heading = ps[vysn_start]._p
    parent = vysn_heading.getparent()

    new_paragraphs = [
        "За час переддипломної практики розроблено веб-платформу TechBox для продажу електроніки та ґаджетів на українському ринку. Робота охопила повний цикл створення програмного продукту: від аналізу предметної області електронної комерції та проєктування архітектури до реалізації, тестування й розгортання готового рішення.",
        "На етапі аналізу досліджено бізнес-процеси інтернет-торгівлі та проблеми існуючих рішень, що дозволило сформувати вимоги до системи. Обрано технологічний стек, який поєднує серверний фреймворк FastAPI, бібліотеку React 19 для клієнтської частини, реляційну СКБД PostgreSQL та контейнеризацію через Docker. Спроєктовано реляційну базу з шести таблиць, що охоплює товари, користувачів, замовлення, відгуки та промокоди.",
        "Серверна частина містить понад двадцять ендпоінтів REST API із JWT-авторизацією та ролевою моделлю доступу. Клієнтський односторінковий застосунок забезпечує перегляд каталогу, роботу з кошиком, оформлення замовлення, особистий кабінет і систему відгуків. Адміністративна панель надає статистику продажів, управління замовленнями, товарами, користувачами та промокодами. Інтегровано LiqPay із підтримкою webhook для автоматичного оновлення статусу оплати, а також AI-модуль генерації описів товарів на базі Google Gemini із кешуванням результатів.",
        "Створена система придатна як основа для реального інтернет-магазину. Інтерфейс та AI-описи працюють українською мовою, платежі здійснюються через LiqPay. Набутий досвід охоплює проєктування REST API, роботу з ORM та міграціями, побудову SPA на React із TypeScript, інтеграцію зовнішніх сервісів (LiqPay, Google Gemini, SMTP), контейнеризацію та оркестрацію."
    ]

    # Insert after vysn_heading
    WN = W[:-1] + '}'  # '{http://...}/main}'

    insert_after = vysn_heading
    for text in new_paragraphs:
        new_p = etree.Element(WN + 'p')
        pPr = etree.SubElement(new_p, WN + 'pPr')
        spacing = etree.SubElement(pPr, WN + 'spacing')
        spacing.set(WN + 'line', '360')
        spacing.set(WN + 'lineRule', 'auto')
        ind = etree.SubElement(pPr, WN + 'ind')
        ind.set(WN + 'firstLine', '567')
        jc = etree.SubElement(pPr, WN + 'jc')
        jc.set(WN + 'val', 'both')

        r = etree.SubElement(new_p, WN + 'r')
        rPr = etree.SubElement(r, WN + 'rPr')
        rFonts = etree.SubElement(rPr, WN + 'rFonts')
        rFonts.set(WN + 'ascii', 'Times New Roman')
        rFonts.set(WN + 'hAnsi', 'Times New Roman')
        rFonts.set(WN + 'cs', 'Times New Roman')
        sz = etree.SubElement(rPr, WN + 'sz')
        sz.set(WN + 'val', '28')
        szCs = etree.SubElement(rPr, WN + 'szCs')
        szCs.set(WN + 'val', '28')

        wt = etree.SubElement(r, WN + 't')
        wt.text = text
        wt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        insert_after.addnext(new_p)
        insert_after = new_p

    print(f"FIX 2: inserted {len(new_paragraphs)} flowing ВИСНОВКИ paragraphs")

# ============================================================
# FIX 3: Dash list items — lowercase first letter + ";" (last gets ".")
# ============================================================
# Re-read paragraphs after modifications
ps = d.paragraphs

# Proper nouns / acronyms that stay uppercase
PROPER = {'Docker', 'Alembic', 'LiqPay', 'Google', 'FastAPI', 'AI', 'PostgreSQL',
          'React', 'JWT', 'REST', 'SPA', 'Python', 'SMTP', 'Gemini', 'API',
          'FastAPI-Mail', 'TypeScript', 'CSS', 'HTML', 'HTTP', 'HTTPS', 'URL', 'URI'}

# Build groups of consecutive numId=29
groups = []
current_group = []
for i, p in enumerate(ps):
    if numid(p) == '29':
        current_group.append(i)
    else:
        if current_group:
            groups.append(current_group)
            current_group = []
if current_group:
    groups.append(current_group)

n_lower = 0
n_semi = 0
for group in groups:
    for idx, pi in enumerate(group):
        p = ps[pi]
        t = p.text.strip()
        if not t:
            continue

        is_last = (idx == len(group) - 1)

        # Fix ending: ";" for all except last which gets "."
        desired_end = '.' if is_last else ';'
        if t[-1] in ('.', ';', ',', ':'):
            if t[-1] != desired_end:
                # Replace in last run
                for r in reversed(p.runs):
                    rt = r.text.rstrip()
                    if rt:
                        if rt[-1] in ('.', ';', ',', ':'):
                            r.text = r.text[:-(len(r.text) - len(r.text.rstrip()))] if r.text != r.text.rstrip() else r.text
                            r.text = r.text[:-1] + desired_end
                        n_semi += 1
                        break
        else:
            # No punctuation — add desired_end
            for r in reversed(p.runs):
                if r.text.strip():
                    r.text = r.text.rstrip() + desired_end
                    n_semi += 1
                    break

        # Fix capitalization: lowercase first letter unless proper noun
        first_word = t.split()[0] if t else ''
        # Check if first word (or its base without trailing punctuation) is proper
        first_clean = re.sub(r'[^a-zA-Zа-яА-ЯіІїЇєЄґҐ\'-]', '', first_word)
        is_proper = first_clean in PROPER or any(first_clean.startswith(pn) for pn in PROPER)

        if not is_proper and t[0].isupper():
            # Lowercase the first character in the first non-empty run
            for r in p.runs:
                rt = r.text
                # Find first actual letter
                for ci, ch in enumerate(rt):
                    if ch.isalpha():
                        r.text = rt[:ci] + ch.lower() + rt[ci+1:]
                        n_lower += 1
                        break
                break

print(f"FIX 3: lowercased {n_lower} first letters, fixed {n_semi} endings")

# ============================================================
# FIX 4: Section 1.1 — expand practice base description
# ============================================================
ps = d.paragraphs
for i, p in enumerate(ps):
    t = p.text.strip()
    # Find the body 1.1 heading (not TOC line)
    if t.startswith('1.1') and '\t' not in p.text and 'Опис бази практики' in t:
        # Next paragraph (i+1) is the short base description
        next_p = ps[i + 1]
        nt = next_p.text.strip()
        if nt.startswith('Практика проводилася на базі ФОП'):
            # Replace this short paragraph with expanded text
            new_text = (
                "Практика проводилася на базі ФОП Раделицький Володимир Степанович "
                "(Львівська обл., Миколаївський р-н, с. Вербіж, вул. Івана Франка, 6/4). "
                "Підприємство спеціалізується на роздрібній та дистанційній торгівлі електронікою, "
                "комп'ютерною периферією та побутовими ґаджетами для українського ринку. "
                "Основним каналом збуту є інтернет-магазин, що забезпечує повний цикл "
                "обслуговування клієнтів: від перегляду каталогу та оформлення замовлення "
                "до онлайн-оплати й доставки. У межах практики перед студентом поставлено "
                "завдання розробити нову версію веб-платформи TechBox, яка замінить попереднє "
                "рішення та підвищить ефективність продажів завдяки сучасному технологічному стеку."
            )
            # Clear existing runs and write new text
            for r in list(next_p._p.findall(W + 'r')):
                next_p._p.remove(r)
            run = next_p.add_run(new_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            print(f"FIX 4: expanded section 1.1 base description at paragraph {i+1}")
            break

d.save(SRC)
print(f"\nSaved: {SRC} ({os.path.getsize(SRC)/1024:.1f} KB)")
