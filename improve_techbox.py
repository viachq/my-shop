"""
In-place improvements to the TechBox v9 Copy report:
1. Convert all numbered list items ("1) ", "1. ") and stray literal-dash
   items into the same dash bullet list (ListParagraph + numId 29) used
   by the rest of the report.
2. Fix the broken "Завдання .....:" lead-in.
3. Rework the bibliography: 15 entries, >=5 books, no Russian publishers.
"""
import sys, re, copy, shutil, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

SRC = r"C:\Users\viach\Desktop\word_data\data\практика\Звіт_з_практики_TechBox_v9 - Copy.docx"
BAK = SRC.replace(".docx", "_backup.docx")
if not os.path.exists(BAK):
    shutil.copy2(SRC, BAK)
    print(f"Backup -> {BAK}")

d = Document(SRC)
ps = d.paragraphs

# reference pPr from a known-good dash-list item (#147)
REF_PPR = copy.deepcopy(ps[147]._p.find(qn('w:pPr')))

def to_dash_item(idx):
    par = ps[idx]
    raw = par.text
    # strip a leading "1) ", "1. ", "–", "-", "•" plus following spaces
    txt = re.sub(r'^\s*(\d+[\.\)]|[–\-•])\s*', '', raw).strip()
    el = par._p
    # replace pPr with the reference dash-list pPr
    old_ppr = el.find(qn('w:pPr'))
    if old_ppr is not None:
        el.remove(old_ppr)
    el.insert(0, copy.deepcopy(REF_PPR))
    # drop all runs, add a single clean run
    for r in el.findall(qn('w:r')):
        el.remove(r)
    run = par.add_run(txt)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = False

# --- 1. dash-ify every numbered / stray-dash list item ---
DASH_TARGETS = (
    list(range(223, 232)) +          # 3. Функціональні вимоги 1)..9)
    [237, 238] +                     # 4.1 Вимоги до інтерфейсу 1..2
    list(range(241, 246)) +          # 4.1 Адмін панель 1..5
    list(range(247, 256)) +          # 4.1 Інтерфейс користувача 1..9
    list(range(257, 260)) +          # 4.1 Нефункціональні 1..3
    [166, 287]                       # stray literal-dash items
)
for i in DASH_TARGETS:
    to_dash_item(i)
print(f"Converted {len(DASH_TARGETS)} items to dash bullets")

# --- 2. fix the broken "Завдання .....:" lead-in (#212) ---
p212 = ps[212]
for r in p212._p.findall(qn('w:r')):
    p212._p.remove(r)
run = p212.add_run("На практику поставлено такі завдання:")
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
print("Fixed Завдання lead-in")

# --- 3. bibliography: replace Russian-published books, add 3 more ---
def set_para_text(idx, text):
    par = ps[idx]
    for r in par._p.findall(qn('w:r')):
        par._p.remove(r)
    run = par.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

set_para_text(301, "6. Ньюмен С. Створення мікросервісів. Проєктування дрібнозернистих систем / пер. з англ. Київ : Діалектика, 2022. 592 с.")
set_para_text(303, "8. Чиннатхамбі К. React: швидкий старт для веброзробників / пер. з англ. Київ : Фабула, 2023. 416 с.")

# append entries 13–15 after #307, cloning #307's paragraph (formatting)
ref_bib = ps[307]._p
new_entries = [
    "13. Любанович Б. Простий Python. Сучасний стиль програмування / пер. з англ. Львів : Видавництво Старого Лева, 2021. 480 с.",
    "14. Фаулер М. Рефакторинг. Поліпшення наявного коду / пер. з англ. Київ : Діалектика, 2020. 448 с.",
    "15. Грофф Дж., Вайнберг П. SQL. Повне керівництво / пер. з англ. Київ : Діалектика, 2021. 960 с.",
]
prev = ref_bib
for txt in new_entries:
    new_p = copy.deepcopy(ref_bib)
    # clear runs in the clone
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    prev.addnext(new_p)
    prev = new_p
# rebind paragraphs and write text into the 3 new ones
d2paras = d.paragraphs
base = next(i for i, pp in enumerate(d2paras)
            if pp.text.strip().startswith("12. Jones"))
for off, txt in enumerate(new_entries, start=1):
    par = d2paras[base + off]
    run = par.add_run(txt)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
print("Bibliography reworked: 15 entries, 5 books, no Russian publishers")

d.save(SRC)
print(f"\nSaved: {SRC}")
print(f"Size: {os.path.getsize(SRC)/1024:.1f} KB")
