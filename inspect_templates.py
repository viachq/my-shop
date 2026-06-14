import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

print("=" * 80)
print("  TITLE TEMPLATE")
print("=" * 80)
doc = Document(r"C:\Users\viach\Desktop\word_data\data\практика\ДР Практика_Титулка (1).docx")
for i, p in enumerate(doc.paragraphs):
    runs_info = []
    for r in p.runs:
        sz = r.font.size
        b = r.bold
        u = r.underline
        runs_info.append(f"[{repr(r.text[:80])} sz={sz} b={b} u={u}]")
    if runs_info:
        print(f"P{i}: {' | '.join(runs_info)}")
    else:
        print(f"P{i}: (empty)")

print("\n" + "=" * 80)
print("  ZAVDANNYA TEMPLATE (first 83 paragraphs)")
print("=" * 80)
doc2 = Document(r"C:\Users\viach\Desktop\word_data\data\практика\ШАБЛОН_Звіту_з_практики_+_зразок_змісту (2).docx")
for i, p in enumerate(doc2.paragraphs):
    if i > 82:
        break
    runs_info = []
    for r in p.runs:
        sz = r.font.size
        b = r.bold
        u = r.underline
        runs_info.append(f"[{repr(r.text[:80])} sz={sz} b={b} u={u}]")
    if runs_info:
        print(f"P{i}: {' | '.join(runs_info)}")
    else:
        print(f"P{i}: (empty)")
