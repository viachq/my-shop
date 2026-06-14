"""Keep figure/table captions on the same page as their object."""
import sys, os, shutil
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

W = qn  # alias
SRC = r"C:\Users\viach\Desktop\word_data\data\практика\Звіт_з_практики_TechBox_v9 - Copy.docx"
BAK = SRC.replace(".docx", "_backup3.docx")
if not os.path.exists(BAK):
    shutil.copy2(SRC, BAK)
    print("backup ->", BAK)

d = Document(SRC)

def has_drawing(p):
    return p._p.find('.//' + qn('w:drawing')) is not None

paras = d.paragraphs

fig_img = 0
tbl_cap = 0
fig_cap = 0
for i, p in enumerate(paras):
    t = p.text.strip()
    # figure: image paragraph must stay with the caption that follows it
    if has_drawing(p):
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True
        fig_img += 1
    # figure caption: keep the caption lines together (no internal split)
    if t.startswith("Рис."):
        p.paragraph_format.keep_together = True
        fig_cap += 1
    # table caption: must stay with the table that follows it
    if t.startswith("Таблиця"):
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.keep_together = True
        tbl_cap += 1

# tables: forbid a single row from splitting across pages
for tbl in d.tables:
    for row in tbl.rows:
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn('w:cantSplit')) is None:
            trPr.append(OxmlElement('w:cantSplit'))

print(f"figure images keep_with_next: {fig_img}")
print(f"figure captions keep_together: {fig_cap}")
print(f"table captions keep_with_next: {tbl_cap}")
print(f"tables row-split disabled: {len(d.tables)} tables")

d.save(SRC)
print("saved:", f"{os.path.getsize(SRC)/1024:.1f} KB")
