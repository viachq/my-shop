import sys, os, copy
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from lxml import etree

TITLE = r"C:\Users\viach\Desktop\word_data\data\практика\ДР Практика_Титулка (1) (2).docx"
ZAVD  = r"C:\Users\viach\Desktop\word_data\data\практика\ШАБЛОН_Звіту_з_практики_+_зразок_змісту.docx"
OUT   = r"C:\Users\viach\Desktop\word_data\data\практика\Титулка_та_завдання_v2.docx"

NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def before_sectpr(body, el):
    s = body.find('w:sectPr', NS)
    if s is not None:
        s.addprevious(el)
    else:
        body.append(el)

# Open title as base - exact 1:1 copy
doc = Document(TITLE)
body = doc.element.body

# Get title section properties (margins, page size)
title_sectpr = body.find('w:sectPr', NS)
title_sp_copy = copy.deepcopy(title_sectpr)

# Remove any footer/header references from title sectPr copy
# (to avoid page numbers leaking through)
for tag in ['w:headerReference', 'w:footerReference']:
    for ref in title_sp_copy.findall(tag, NS):
        title_sp_copy.remove(ref)

# Insert section break paragraph after title content
# This preserves title page layout in its own section
break_p = etree.Element(f'{{{W}}}p')
break_ppr = etree.SubElement(break_p, f'{{{W}}}pPr')
break_ppr.append(title_sp_copy)
st = etree.SubElement(title_sp_copy, f'{{{W}}}type')
st.set(f'{{{W}}}val', 'nextPage')
before_sectpr(body, break_p)

# Open zavdannya template
doc_z = Document(ZAVD)

# Copy P0 through P50 (first 2 pages; page break is IN P50)
# But we need to REMOVE the page break from P50 since it's the last thing we copy
p_count = 0
for child in list(doc_z.element.body):
    tag = etree.QName(child.tag).localname
    if tag == 'sectPr':
        continue
    cloned = copy.deepcopy(child)
    if tag == 'p':
        if p_count > 82:
            break
        # If this is P82, remove the page break element from it (end of page 2)
        if p_count == 82:
            for br in cloned.findall('.//w:br', NS):
                btype = br.get(f'{{{W}}}type')
                if btype == 'page':
                    br.getparent().remove(br)
        before_sectpr(body, cloned)
        p_count += 1
    elif tag == 'tbl':
        before_sectpr(body, cloned)

print(f"Copied {p_count} paragraphs from Завдання (first 2 pages)")

# Replace body's sectPr with zavdannya's section properties (margins/page size)
zav_sectpr = doc_z.element.body.find('w:sectPr', NS)
if zav_sectpr is not None:
    new_sp = copy.deepcopy(zav_sectpr)
    # Remove header/footer refs to avoid broken relationships
    for tag in ['w:headerReference', 'w:footerReference']:
        for ref in new_sp.findall(tag, NS):
            new_sp.remove(ref)
    old_sp = body.find('w:sectPr', NS)
    if old_sp is not None:
        body.replace(old_sp, new_sp)

# Clear ALL footers in all sections - NO page numbering
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        # Remove all child elements
        for elem in list(p._element):
            p._element.remove(elem)

print("Cleared all footers - no page numbering")

doc.save(OUT)
print(f"\nSaved: {OUT}")
print(f"Size: {os.path.getsize(OUT)/1024:.1f} KB")
