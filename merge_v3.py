import sys, os, copy
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docxcompose.composer import Composer
from lxml import etree

TITLE = r"C:\Users\viach\Desktop\word_data\data\практика\ДР Практика_Титулка (1) (2).docx"
ZAVD  = r"C:\Users\viach\Desktop\word_data\data\практика\ШАБЛОН_Звіту_з_практики_+_зразок_змісту.docx"
TMP   = r"C:\Users\viach\Desktop\word_data\data\практика\_zavd_2pages_tmp.docx"
OUT   = r"C:\Users\viach\Desktop\word_data\data\практика\Титулка_та_завдання_v3.docx"

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = {'w': W}

# --- Step 1: build truncated zavdannya = first 2 pages only (P0..P82) ---
zdoc = Document(ZAVD)
zbody = zdoc.element.body

# collect paragraph elements in order
p_idx = 0
to_remove = []
for child in list(zbody):
    tag = etree.QName(child.tag).localname
    if tag == 'p':
        if p_idx == 82:
            # remove the page break inside P82 (it's end of page 2)
            for br in child.findall('.//w:br', NS):
                if br.get(f'{{{W}}}type') == 'page':
                    br.getparent().remove(br)
        elif p_idx > 82:
            to_remove.append(child)
        p_idx += 1
    elif tag == 'tbl':
        # keep tables only if they appear before we pass P82
        if p_idx > 82:
            to_remove.append(child)
    elif tag == 'sectPr':
        continue  # keep document-final sectPr

for el in to_remove:
    zbody.remove(el)

zdoc.save(TMP)
print(f"Truncated zavdannya saved ({p_idx if p_idx<=82 else 83} paras kept up to P82)")

# --- Step 2: compose: title (master) + truncated zavdannya ---
master = Document(TITLE)
composer = Composer(master)
zpart = Document(TMP)
composer.append(zpart)  # docxcompose inserts a page/section break and merges styles
composer.save(OUT)

# --- Step 3: remove ALL footers/page numbering from every section ---
out = Document(OUT)
for section in out.sections:
    for which in (section.footer, section.first_page_footer, section.even_page_footer):
        try:
            which.is_linked_to_previous = False
        except Exception:
            pass
        for p in which.paragraphs:
            for elem in list(p._element):
                p._element.remove(elem)
out.save(OUT)

os.remove(TMP)
print(f"\nSaved: {OUT}")
print(f"Size: {os.path.getsize(OUT)/1024:.1f} KB")
