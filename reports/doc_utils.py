"""
Shared formatting utilities for generating academic .docx reports.
Ukrainian academic standards (DSTU 3008:2015).
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCREENSHOTS_DIR = os.path.join(os.path.dirname(__file__), 'screenshots')


def create_document():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.color.rgb = RGBColor(0, 0, 0)

    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)

    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)

    rpr = style.element.get_or_add_rPr()
    rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Times New Roman" w:cs="Times New Roman"/>')
    rpr.append(rfonts)

    add_page_numbers(doc)

    return doc


def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)

        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)

        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)

        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run3._r.append(fldChar2)

        run4 = p.add_run('1')
        run4.font.size = Pt(12)
        run4.font.name = 'Times New Roman'

        run5 = p.add_run()
        fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run5._r.append(fldChar3)


def add_empty_paragraph(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run('')
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'


def add_centered_text(doc, text, bold=False, size=14, spacing_before=0, spacing_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_right_text(doc, text, bold=False, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_left_text(doc, text, bold=False, size=14, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    return p


def add_body_mixed(doc, parts):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    for text, bold in parts:
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run.bold = bold
    return p


def add_inline_section_heading(doc, number, title):
    text = f"{number}. {title.upper()}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True

    p.style = doc.styles['Heading 1']
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(24)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.5


def add_section_heading(doc, number, title):
    text = f"РОЗДІЛ {number}\n{title.upper()}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True

    p.style = doc.styles['Heading 1']
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(24)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.5


def add_subsection_heading(doc, number, title):
    text = f"{number} {title}" if number else title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True

    p.style = doc.styles['Heading 2']
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(18)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5


def add_unnumbered_heading(doc, title, center=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True

    p.style = doc.styles['Heading 1']
    p.clear()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(24)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.5


def add_dash_list(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f"– {item}")
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'


def add_numbered_list(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f"{i}) {item}")
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'


def add_table(doc, caption, headers, rows):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(caption)
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run.italic = True

    ncols = len(headers)
    nrows = len(rows) + 1
    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(str(val))
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'

    add_empty_paragraph(doc)
    return table


def add_image(doc, filename, caption, width_cm=15):
    img_path = os.path.join(SCREENSHOTS_DIR, filename)
    if not os.path.exists(img_path):
        add_figure_placeholder(doc, caption)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))

    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.first_line_indent = Cm(0)
    pc.paragraph_format.space_after = Pt(12)
    run = pc.add_run(caption)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.italic = True


def add_figure_placeholder(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('[Рисунок]')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(128, 128, 128)

    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.first_line_indent = Cm(0)
    pc.paragraph_format.space_after = Pt(12)
    run = pc.add_run(caption)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.italic = True


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('ЗМІСТ')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = paragraph.add_run('[Ctrl+A, F9]')
    run4.font.size = Pt(12)
    run4.font.name = 'Times New Roman'
    run4.font.color.rgb = RGBColor(128, 128, 128)

    run5 = paragraph.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)

    doc.add_page_break()


def add_page_break(doc):
    doc.add_page_break()


def add_references(doc, refs):
    add_unnumbered_heading(doc, 'Список використаних джерел')
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.hanging_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f"{i}. {ref}")
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
