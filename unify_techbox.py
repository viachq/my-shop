"""Uniformity pass over the TechBox report (in place)."""
import sys, re, os, shutil
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
SRC = r"C:\Users\viach\Desktop\word_data\data\практика\Звіт_з_практики_TechBox_v9 - Copy.docx"
BAK = SRC.replace(".docx", "_backup2.docx")
if not os.path.exists(BAK):
    shutil.copy2(SRC, BAK)
    print("backup ->", BAK)

d = Document(SRC)

def is_pagebreak(p):
    return any(br.get(W + 'type') == 'page' for br in p._p.iter(W + 'br'))

def is_img(p):
    return p._p.find('.//' + W + 'drawing') is not None

def numpr_id(p):
    ppr = p._p.find(W + 'pPr')
    if ppr is None:
        return None
    npr = ppr.find(W + 'numPr')
    if npr is None:
        return None
    nid = npr.find(W + 'numId')
    return nid if nid is not None else None

SUBHEAD = re.compile(r'^\d+\.\d+\s')

paras = d.paragraphs

# ---- 1. unify list numId (28 -> 29) ----
n_num = 0
for p in paras:
    nid = numpr_id(p)
    if nid is not None and nid.get(W + 'val') == '28':
        nid.set(W + 'val', '29')
        n_num += 1
print(f"list numId 28->29: {n_num}")

# ---- 2. body lead-in indent + punctuation + subheading spacing ----
for p in paras:
    t = p.text.strip()
    if t == 'Функціональні вимоги:' and numpr_id(p) is None:
        p.paragraph_format.first_line_indent = Cm(1)
        print("fixed indent: 'Функціональні вимоги:'")
    if t == 'Функціональні вимоги.':
        for r in p.runs:
            if r.text.strip():
                r.text = r.text.replace('Функціональні вимоги.', 'Функціональні вимоги:')
        print("fixed punctuation: 'Функціональні вимоги.' -> ':'")
    if t.startswith('2.2 Огляд сучасних засобів розробки'):
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        print("fixed subheading 2.2 spacing 12/6")

# ---- 3. table captions: drop dash to match de-dashed figure captions ----
cap_re = re.compile(r'^(Таблиця\s+\d+\.\d+)\s*[–—-]\s*')
for p in paras:
    t = p.text.strip()
    if t.startswith('Таблиця') and cap_re.match(t):
        new = cap_re.sub(r'\1 ', t)
        # rewrite as single run, keep caption font (italic 12 TNR)
        runs = p.runs
        font0 = runs[0].font if runs else None
        for r in list(p._p.findall(W + 'r')):
            p._p.remove(r)
        run = p.add_run(new)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.italic = True
        print(f"de-dashed caption: {new[:40]}")

# ---- 4. remove inconsistent empty separator paragraphs ----
# rule: an EMPTY (no text / no pagebreak / no image / no numPr) paragraph that
# immediately precedes a subsection heading, a 'Таблиця' caption, or an image
# paragraph -> delete (spacing is handled by paragraph spacing instead).
paras = d.paragraphs  # refresh
def empty(p):
    return (not p.text.strip()) and (not is_pagebreak(p)) and (not is_img(p)) \
           and numpr_id(p) is None

to_del = []
for i, p in enumerate(paras):
    if not empty(p):
        continue
    # find next non-empty paragraph
    j = i + 1
    while j < len(paras) and empty(paras[j]):
        j += 1
    if j >= len(paras):
        continue
    nxt = paras[j]
    nt = nxt.text.strip()
    if SUBHEAD.match(nt) or nt.startswith('Таблиця') or is_img(nxt):
        to_del.append(p)

for p in to_del:
    p._p.getparent().remove(p._p)
print(f"removed {len(to_del)} inconsistent empty paragraphs")

d.save(SRC)
print("saved:", SRC, f"{os.path.getsize(SRC)/1024:.1f} KB")
