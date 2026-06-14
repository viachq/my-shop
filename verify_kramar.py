import sys, zipfile, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

p = r"C:\Users\viach\Desktop\word_data\data\практика\Звіт_з_практики_Kramar.docx"
d = Document(p)
body = d.element.body
ch = list(body)
print("last child:", ch[-1].tag.split('}')[-1])
print("sections:", len(d.sections), "paras:", len(d.paragraphs), "tables:", len(d.tables))

# media in zip
z = zipfile.ZipFile(p)
media = [n for n in z.namelist() if n.startswith('word/media/')]
print("media files:", media)

# which image rIds are referenced in document.xml
docxml = z.read('word/document.xml').decode('utf-8', 'ignore')
embeds = set(re.findall(r'r:embed="([^"]+)"', docxml))
print("referenced embeds:", sorted(embeds))

rels = z.read('word/_rels/document.xml.rels').decode('utf-8', 'ignore')
img_rels = re.findall(r'<Relationship Id="([^"]+)"[^>]*image[^>]*Target="([^"]+)"', rels)
print("image relationships:")
for rid, tgt in img_rels:
    used = "USED" if rid in embeds else "ORPHAN"
    print(f"  {rid} -> {tgt}  [{used}]")
